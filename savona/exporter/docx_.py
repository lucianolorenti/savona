import base64
import json
import os
import tempfile
from io import BytesIO

import mistletoe
import nbformat
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from mistletoe.ast_renderer import ASTRenderer
from PIL import Image
from savona.exporter import Exporter

HEADING_TYPE = 'Heading'
RAWTEXT_TYPE = 'RawText'
PARAGRAPH_TYPE = 'Paragraph'
LIST_TYPE = 'List'
LISTITME_TYPE = 'ListItem'


STREM_OUTPUT_TYPE = 'stream'

IMAGE_OUTPUT_TYPE = 'image/png'
TEXT_OUTPUT_TYPE = 'text/plain'


def table_from_pandas(document, data: pd.DataFrame, include_header=True):
    nrows = data.shape[0]
    if include_header:
        nrows += 1
    table = document.add_table(rows=nrows, cols=data.shape[1])
    drow = 0
    if include_header:
        for i, column in enumerate(data):
            table.cell(0, i).text = data.columns[i]
            drow = 1
    for i, column in enumerate(data):

        for row in range(data.shape[0]):
            table.cell(row+drow, i).text = str(data[column][data.index[row]])


def add_image_from_figure(doc, fig):
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, 'tmp_figure.png')
        fig.savefig(path)
        doc.add_picture(
            path,
            width=effective_width(doc.sections[0]))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def effective_width(section):
    return section.page_width - section.bottom_margin - section.left_margin


class DocxExporter(Exporter):
    def __init__(self, config={}):
        self.doc = Document()
        self.config = config

    @property
    def header(self):
        return self.doc.sections[0].header

    @property
    def footer(self):
        return self.doc.sections[0].footer

    def export(self, notebook_path, output_path):
        with open(notebook_path, 'r') as file:
            content = file.read()
        notebook = nbformat.reads(content, as_version=4)

        self.footer.add_paragraph('HOLA!')

        for node in notebook.cells:
            if node['cell_type'] == 'code':
                self.add_output(node['outputs'])
            if node['cell_type'] == 'markdown':
                self.add_content(node)

        self.doc.save('demo.docx')

    def add_output(self, cell_code):
        for o in cell_code:

            if o['output_type'] == STREM_OUTPUT_TYPE:
                self.doc.add_paragraph(o['text'])
                continue
            data = o['data']
            data_type = list(data.keys())[0]

            if data_type == TEXT_OUTPUT_TYPE:
                self.doc.add_paragraph(data[data_type])
            elif data_type == IMAGE_OUTPUT_TYPE:

                fp = tempfile.TemporaryFile(suffix='.png')
                fp.write(base64.b64decode(data[data_type]))
                fp.seek(0)
                self.doc.add_picture(
                    fp, width=effective_width(self.doc.sections[0]))
                self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def flatten(self, elements):
        content = []
        for el in elements:
            if el['type'] == RAWTEXT_TYPE:
                content.append(el['content'])
            elif el['type'] == PARAGRAPH_TYPE:
                content.append(self.flatten(el['children']))
        return ' '.join(content)

    def add_list(self, node):
        for el in node['children']:
            if el['type'] == LISTITME_TYPE:
                self.doc.add_paragraph(self.flatten(el['children']),
                                       style='List Bullet')

    def process_node(self, node):

        if node['type'] == HEADING_TYPE:
            self.doc.add_heading(
                self.flatten(node['children']),
                level=node['level'])
        elif node['type'] == PARAGRAPH_TYPE:
            self.doc.add_paragraph(self.flatten(node['children']))
        elif node['type'] == LIST_TYPE:
            self.add_list(node)

    def process_elements(self, elements):
        for el in elements:
            self.process_node(el)

    def add_content(self, node):
        elements = json.loads(mistletoe.markdown(node['source'], ASTRenderer))

        self.process_elements(elements['children'])
